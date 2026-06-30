import sys
import io

class DummyWriter:
    def write(self, s):
        pass
    def flush(self):
        pass

if sys.stdout is None:
    sys.stdout = DummyWriter()
if sys.stderr is None:
    sys.stderr = DummyWriter()

if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except:
        pass

import os
import json
import io
import base64
import telebot
import fitz  
import re
import shutil
import math
import smtplib
import imaplib
import email
import mimetypes
import httpx  
import traceback
import time
import threading
import platform
import socket
import logging
from datetime import datetime, timedelta
from email.message import EmailMessage
from email.header import decode_header, Header
from PIL import Image, ImageDraw
from openai import OpenAI  
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from telebot import types, apihelper 
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, BackgroundTasks, Request, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uuid

# =====================================================================
# ЗАГРУЗКА ПЕРЕМЕННЫХ ОКРУЖЕНИЯ ИЗ .env
# =====================================================================
load_dotenv()

# =====================================================================
# 🛠 КОРРЕКТНЫЙ РАСЧЕТ ПУТЕЙ ДЛЯ PYINSTALLER (.EXE РЕЖИМ)
# =====================================================================
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TEMPLATE_FOLDER = os.path.join(BASE_DIR, "Templates")
STATIC_ATTACHMENTS_FOLDER = os.path.join(BASE_DIR, "static_attachments")
COUNTERS_FILE = os.path.join(BASE_DIR, "history_counters.json")
ARCHIVE_FOLDER = os.path.join(BASE_DIR, "FINAL_ARCHIVE")
LOG_FILE = os.path.join(BASE_DIR, "terminator_errors.log")

# =====================================================================
# URL WEB APP (задайте в .env переменную WEBAPP_URL)
# =====================================================================
WEBAPP_URL = os.getenv("WEBAPP_URL", "https://your-domain.com/webapp/index.html")

# =====================================================================
# ПЕРЕХВАТ ВНУТРЕННИХ ОШИБОК БИБЛИОТЕКИ TELEBOT
# =====================================================================
logging.basicConfig(filename=LOG_FILE, level=logging.ERROR, 
                    format='[%(asctime)s] ❌ ВНУТРЕННЯЯ ОШИБКА БИБЛИОТЕКИ: %(message)s')

# =====================================================================
# ФУНКЦИЯ ЛОГГЕРА ДЛЯ ОТСЛЕЖИВАНИЯ СБОЕВ В НАШЕМ КОДЕ
# =====================================================================
def log_error(context_message, error_exception=None):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_msg = f"[{timestamp}] ❌ КРИТИЧЕСКИЙ СБОЙ: {context_message}\n"
    if error_exception:
        log_msg += f"Детали ошибки: {str(error_exception)}\n"
        log_msg += f"Стек вызовов:\n{traceback.format_exc()}\n"
    log_msg += "-"*80 + "\n"
    
    # Выводим в stderr, чтобы логи были видны в панели управления Render
    try:
        sys.stderr.write(log_msg)
        sys.stderr.flush()
    except:
        pass
        
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(log_msg)
    except:
        pass

# =====================================================================
# ГЛОБАЛЬНЫЙ ПЕРЕХВАТ НЕОБРАБОТАННЫХ ИСКЛЮЧЕНИЙ (УЛЬТРА-ЗАЩИТА)
# =====================================================================
def handle_unhandled_exception(exc_type, exc_value, exc_traceback):
    """Ловит любые падения главного потока вне блоков try-except"""
    if issubclass(exc_type, KeyboardInterrupt):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    log_error("Необработанное исключение системы (Критическое падение скрипта)", exc_value)

sys.excepthook = handle_unhandled_exception

def handle_thread_exception(args):
    """Ловит скрытые падения внутри фоновых потоков (демонов почты и трея)"""
    log_error(f"Необработанное исключение в фоновом потоке [{args.thread.name}]", args.exc_value)

threading.excepthook = handle_thread_exception

# =====================================================================
# НАСТРОЙКИ КОНВЕЙЕРА И ТОКЕНЫ (из .env)
# =====================================================================
OPENROUTER_API_KEY  = os.getenv("OPENROUTER_API_KEY", "")
TELEGRAM_BOT_TOKEN  = os.getenv("TELEGRAM_BOT_TOKEN", "")
MY_CHAT_ID          = int(os.getenv("MY_CHAT_ID", "0"))

SECOND_BOT_TOKEN    = os.getenv("SECOND_BOT_TOKEN", "")
SECOND_BOT_CHAT_ID  = int(os.getenv("SECOND_BOT_CHAT_ID", "0"))

CURRENT_MODEL   = "google/gemini-3.5-flash"
BOT_START_TIME  = datetime.now()

SMTP_SERVER  = os.getenv("SMTP_SERVER", "smtp.mail.ru")
SMTP_PORT    = int(os.getenv("SMTP_PORT", "465"))
IMAP_SERVER  = os.getenv("IMAP_SERVER", "imap.mail.ru")
IMAP_PORT    = int(os.getenv("IMAP_PORT", "993"))
EMAIL_SENDER   = os.getenv("EMAIL_SENDER", "")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD", "")

EMAIL_OPTIONS = [
    "aquatechexpert@mail.ru",
    "agent@zddrostov.com",
    "agent@seagatenvr.ru",
    "maritimenvrsk@mail.ru"
]

PDF_START_FROM_FILE = 4

COMPANIES_DATA = {
    "almar": {
        "title": "ООО «Альмар Сервисиз»",
        "company_contract": "Договор с ООО «Альмар Сервисиз» № 5 от 01.06.2026/Agreement with «Almar Services» LTD No.5 dd 01.06.2026",
        "company_short": "Альмар Сервисиз",
        "company_full": "ООО «Альмар Сервисиз»/Almar Services LTD."
    },
    "zheldor": {
        "title": "ООО «ЖелдорДоставка»",
        "company_contract": "Договор с ООО «ЖелдорДоставка» № 4 от 20.04.2026/Agreement with «ZheldorDostavka» LTD No.4 dd 20.04.2026",
        "company_short": "ЖелдорДоставка",
        "company_full": "ООО «ЖелдорДоставка»/ZheldorDostavka LTD."
    },
    "seagate": {
        "title": "ООО «МА СИГЕЙТ»",
        "company_contract": "Договор с ООО «МА СИГЕЙТ» № 1 от 15.04.2026/Agreement with «MA SEAGATE» LTD No.1 dd 15.04.2026",
        "company_short": "МА СИГЕЙТ",
        "company_full": "ООО «МА СИГЕЙТ»/MA SEAGATE LTD."
    }
}

PROXY_URL = os.getenv("PROXY_URL", "").strip()
BOT_ACTIVE = True

if PROXY_URL:
    http_client = httpx.Client(proxy=PROXY_URL, timeout=60.0)
    apihelper.proxy = {'http': PROXY_URL, 'https': PROXY_URL}
else:
    http_client = httpx.Client(timeout=60.0)

client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY, http_client=http_client)
bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN)

# Настройка тайм-аутов для предотвращения разрыва соединения при отправке тяжелых архивов
apihelper.CONNECT_TIMEOUT = 60
apihelper.READ_TIMEOUT = 120

second_bot = None
if SECOND_BOT_TOKEN:
    try:
        second_bot = telebot.TeleBot(SECOND_BOT_TOKEN)
    except:
        pass

user_states = {}

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.heic', '.svg'}

def is_image_file(filename):
    return os.path.splitext(filename.lower())[1] in IMAGE_EXTENSIONS

# =====================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ МОНИТОРИНГА
# =====================================================================
def get_dir_size_mb(folder_path):
    if not os.path.exists(folder_path):
        return 0.0
    total_size = 0
    for dirpath, dirnames, filenames in os.walk(folder_path):
        for f in filenames:
            fp = os.path.join(dirpath, f)
            if os.path.exists(fp):
                total_size += os.path.getsize(fp)
    return round(total_size / (1024 * 1024), 2)

def get_today_acts_count(db_data):
    acts = db_data.get("used_acts", [])
    today_str = datetime.now().strftime("%d.%m.%Y")
    return sum(1 for a in acts if str(a.get("timestamp", "")).startswith(today_str))

def count_templates():
    if not os.path.exists(TEMPLATE_FOLDER):
        return 0
    return len([f for f in os.listdir(TEMPLATE_FOLDER) if f.lower().endswith('.docx')])

def compile_advanced_status_text():
    uptime = datetime.now() - BOT_START_TIME
    db = load_history_db()
    archive_size = get_dir_size_mb(ARCHIVE_FOLDER)
    today_count = get_today_acts_count(db)
    templates_count = count_templates()
    hostname = socket.gethostname()
    os_platform = platform.system() + " " + platform.release()
    
    return (
        "🤖 **РАСШИРЕННЫЙ СТАТУС КОНВЕЙЕРА «ТЕРМИНАТОР»**\n\n"
        "🟢 **Ядро процесса:**\n"
        f"• Режим работы: `Изолированные Сессии v3.9` (Глобальный Логгер)\n"
        f"• Время непрерывной работы: `{uptime.days}д {uptime.seconds // 3600}ч {(uptime.seconds % 3600) // 60}м`\n"
        f"• ИИ Модель парсера: `{CURRENT_MODEL}`\n\n"
        "📊 **Статистика и Аналитика БД:**\n"
        f"• Текущий глобальный номер (БД): `{db.get('last_doc_number', 0)}`\n"
        f"• Всего актов в истории: `{len(db.get('used_acts', []))}`\n"
        f"• Обработано судов за сегодня: *{today_count}*\n"
        f"• Вес архива документов: `{archive_size} МБ`\n\n"
        "📁 **Конфигурация ресурсов папок:**\n"
        f"• Шаблонов актов (.docx) в работе: `{templates_count} шт`\n"
        f"• Конвертер таблиц: `Движок MS Office Word (Многопоточный СОМ)`\n"
        f"• Задействовано фирм-контрагентов: `{len(COMPANIES_DATA)}`\n\n"
        "🖥 **Служебная телеметрия сервера:**\n"
        f"• Имя рабочего ПК: `{hostname}`\n"
        f"• Операционная система: `{os_platform}`\n"
        f"• Файл логов сбоев: `terminator_errors.log`\n\n"
        f"_Срез данных обновлен: {datetime.now().strftime('%H:%M:%S')}_"
    )

# =====================================================================
# БАЗА ДАННЫХ И УЧЕТ АКТОВ
# =====================================================================
def load_history_db():
    if not os.path.exists(COUNTERS_FILE):
        return {"last_doc_number": 145, "used_acts": []}
    try:
        with open(COUNTERS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data.get("last_doc_number"), int):
                data["last_doc_number"] = 145
            return data
    except Exception as e:
        log_error("Ошибка при чтении history_counters.json", e)
        return {"last_doc_number": 145, "used_acts": []}

def save_history_db(db_data):
    try:
        with open(COUNTERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(db_data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        log_error("Ошибка при сохранении истории в БД", e)

def get_next_act_number_for_date(target_date_str):
    try:
        dt = datetime.strptime(target_date_str.strip(), "%d.%m.%Y")
        prefix = f"{dt.strftime('%d')}/{dt.strftime('%m')}/"
    except:
        prefix = f"{datetime.now().strftime('%d/%m/')}"
    db = load_history_db()
    count = sum(1 for a in db.get("used_acts", []) if str(a.get("diving_date", "")).strip() == target_date_str.strip()) + 1
    suffix = f"0{count}" if count < 10 else str(count)
    return f"{prefix}{suffix}"

def log_inspection_event(doc_number, act_number, vessel_name, diving_date, vessel_data=None):
    db = load_history_db()
    try:
        db["last_doc_number"] = int(doc_number)
    except:
        db["last_doc_number"] = 145
    act_entry = {
        "act_number": str(act_number),
        "vessel_name": str(vessel_name),
        "diving_date": str(diving_date),
        "timestamp": datetime.now().strftime("%d.%m.%Y %H:%M"),
        "vessel_data": dict(vessel_data) if vessel_data else {}
    }
    # Удаляем старую запись с таким же номером акта, чтобы перезаписать её актуальными данными
    db["used_acts"] = [a for a in db.get("used_acts", []) if a.get("act_number") != str(act_number)]
    db["used_acts"].append(act_entry)
    save_history_db(db)

# =====================================================================
# РАСЧЕТЫ И ИНСТРУМЕНТЫ СЛИЯНИЯ
# =====================================================================
def extract_json_from_text(text):
    match = re.search(r'\{.*\}', text, re.DOTALL)
    return match.group(0) if match else "{}"

def calculate_underwater_areas(loa, breadth, draft_fore, draft_aft, ship_type):
    try:
        loa_str = str(loa).strip()
        b_str   = str(breadth).strip()
        if loa_str.lower() == 'none' or not loa_str: loa_str = "0.0"
        if b_str.lower()   == 'none' or not b_str:   b_str   = "0.0"
        loa_val  = float(loa_str)
        b_val    = float(b_str)
        df_fore  = float(str(draft_fore).replace(',', '.').strip() or "0.0")
        df_aft   = float(str(draft_aft).replace(',', '.').strip() or "0.0")
        if loa_val == 0.0 or b_val == 0.0:
            return 0.0, 0.0
        lbp   = loa_val * 0.95
        t_avg = (df_fore + df_aft) / 2.0
        st    = str(ship_type).lower()
        cb    = 0.81 if "tanker" in st else (0.80 if "bulk" in st else (0.65 if "container" in st else 0.75))
        wsa   = 1.02 * lbp * (cb * b_val + 1.7 * t_avg)
        if "tanker" in st or "bulk" in st:
            propeller_area = 103.0 if loa_val > 300 else (60.0 if loa_val > 200 else 26.0)
        elif "container" in st:
            propeller_area = 52.9 if loa_val > 180 else 35.0
        else:
            propeller_area = 13.0
        return round(wsa, 2), round(propeller_area, 2)
    except Exception as e:
        log_error("Ошибка при математическом расчете площадей WSA", e)
        return 0.0, 0.0

def update_draft_calculations(vessel_data, draft_fore, draft_aft):
    try:
        vessel_data['draft_fore'] = float(str(draft_fore).replace(',', '.').strip() or "0.0")
        vessel_data['draft_aft']  = float(str(draft_aft).replace(',', '.').strip() or "0.0")
    except:
        vessel_data['draft_fore'], vessel_data['draft_aft'] = 0.0, 0.0
    wsa, prop = calculate_underwater_areas(
        vessel_data.get('loa', 0.0), vessel_data.get('breadth', 0.0),
        vessel_data['draft_fore'], vessel_data['draft_aft'],
        vessel_data.get('ship_type', 'GENERAL CARGO')
    )
    vessel_data.update({'wetted_surface': wsa, 'propeller_area': prop})

def extract_base64_images_from_pdf(pdf_path):
    try:
        fitz.Tools().redirect_stdout(io.StringIO())
    except:
        pass
    try:
        doc = fitz.open(pdf_path)
        base64_images = []
        for i in range(min(len(doc), 1)):
            page = doc.load_page(i)
            pix  = page.get_pixmap(dpi=100)
            img  = Image.open(io.BytesIO(pix.tobytes("jpeg")))
            buf  = io.BytesIO()
            img.save(buf, format="JPEG", quality=75)
            base64_images.append(base64.b64encode(buf.getvalue()).decode('utf-8'))
        doc.close()
        return base64_images
    except Exception as e:
        log_error(f"Не удалось вытащить картинку из PDF спецификации {pdf_path}", e)
        return []

def clean_and_replace_text_nodes(paragraphs, row_dict):
    for p in paragraphs:
        p_text = p.text
        if "{{" not in p_text or "}}" not in p_text:
            continue
        for r in p.runs:
            for key, value in row_dict.items():
                tag = f"{{{{{key}}}}}"
                if tag in r.text:
                    r.text = r.text.replace(tag, str(value))
        still_has_tags = any(
            f"{{{{{k}}}}}" in "".join([r.text for r in p.runs]) for k in row_dict.keys()
        )
        if still_has_tags and len(p.runs) > 1:
            full_text = "".join([r.text for r in p.runs])
            for key, value in row_dict.items():
                tag = f"{{{{{key}}}}}"
                if tag in full_text:
                    full_text = full_text.replace(tag, str(value))
            p.runs[0].text = full_text
            for r in p.runs[1:]:
                r.text = ""

def global_tag_replacer_mono(doc, row_dict):
    clean_and_replace_text_nodes(doc.paragraphs, row_dict)
    for section in doc.sections:
        for h_f in [section.header, section.footer]:
            clean_and_replace_text_nodes(h_f.paragraphs, row_dict)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                clean_and_replace_text_nodes(cell.paragraphs, row_dict)
    for node in doc._element.xpath('.//w:t'):
        if node.text:
            text_node = node.text
            for key, value in row_dict.items():
                tag = f"{{{{{key}}}}}"
                if tag in text_node:
                    text_node = text_node.replace(tag, str(value))
            node.text = text_node

def run_isolated_conversion(docx_path, pdf_path):
    if platform.system() == "Windows":
        try:
            os.system("taskkill /f /im WINWORD.EXE >nul 2>&1")
            time.sleep(0.5)
        except:
            pass
        initialized = False
        try:
            import pythoncom
            pythoncom.CoInitialize()
            initialized = True
        except Exception as e:
            log_error("Не удалось вызвать CoInitialize для потока конвертации", e)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
        except Exception as e:
            log_error(f"Сбой конвертации docx2pdf для файла: {docx_path}", e)
        finally:
            if initialized:
                try:
                    import pythoncom
                    pythoncom.CoUninitialize()
                except:
                    pass
            try:
                os.system("taskkill /f /im WINWORD.EXE >nul 2>&1")
            except:
                pass
    else:
        # Linux / Render: use headless LibreOffice
        import subprocess
        try:
            outdir = os.path.dirname(pdf_path)
            cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path]
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60)
            if result.returncode != 0:
                log_error(f"Сбой LibreOffice ({result.returncode}): {result.stderr.decode('utf-8', errors='ignore')}", None)
        except Exception as e:
            log_error(f"Сбой конвертации через LibreOffice для файла: {docx_path}", e)

# =====================================================================
# ОТПРАВКА ПОЧТЫ (С АВТОСОХРАНЕНИЕМ ЧЕРЕЗ BCC КОПИЮ)
# =====================================================================
import urllib.parse
import ssl

def connect_ssl_via_proxy(target_host, target_port, proxy_url, timeout=120):
    parsed = urllib.parse.urlparse(proxy_url)
    proxy_host = parsed.hostname
    proxy_port = parsed.port or 8000
    
    # Создаем базовый сокет
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.settimeout(timeout)
    sock.connect((proxy_host, proxy_port))
    
    # Формируем заголовки авторизации прокси
    connect_req = f"CONNECT {target_host}:{target_port} HTTP/1.1\r\n"
    connect_req += f"Host: {target_host}:{target_port}\r\n"
    if parsed.username and parsed.password:
        auth_str = f"{parsed.username}:{parsed.password}"
        auth_b64 = base64.b64encode(auth_str.encode('utf-8')).decode('utf-8')
        connect_req += f"Proxy-Authorization: Basic {auth_b64}\r\n"
    connect_req += "\r\n"
    
    sock.sendall(connect_req.encode('utf-8'))
    
    # Читаем статус-ответ прокси
    resp = b""
    while b"\r\n\r\n" not in resp:
        chunk = sock.recv(1024)
        if not chunk:
            break
        resp += chunk
        
    status_line = resp.split(b"\r\n")[0].decode('utf-8', errors='ignore')
    if "200" not in status_line:
        sock.close()
        raise Exception(f"Прокси отказал в CONNECT туннеле: {status_line}")
        
    # Оборачиваем сокет в SSL
    context = ssl.create_default_context()
    ssl_sock = context.wrap_socket(sock, server_hostname=target_host)
    return ssl_sock

def connect_plain_socket_via_proxy(target_host, target_port, proxy_url, timeout=120):
    parsed = urllib.parse.urlparse(proxy_url)
    proxy_host = parsed.hostname
    proxy_port = parsed.port or 8000
    
    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.settimeout(timeout)
    sock.connect((proxy_host, proxy_port))
    
    connect_req = f"CONNECT {target_host}:{target_port} HTTP/1.1\r\n"
    connect_req += f"Host: {target_host}:{target_port}\r\n"
    if parsed.username and parsed.password:
        auth_str = f"{parsed.username}:{parsed.password}"
        auth_b64 = base64.b64encode(auth_str.encode('utf-8')).decode('utf-8')
        connect_req += f"Proxy-Authorization: Basic {auth_b64}\r\n"
    connect_req += "\r\n"
    
    sock.sendall(connect_req.encode('utf-8'))
    
    resp = b""
    while b"\r\n\r\n" not in resp:
        chunk = sock.recv(1024)
        if not chunk:
            break
        resp += chunk
        
    status_line = resp.split(b"\r\n")[0].decode('utf-8', errors='ignore')
    if "200" not in status_line:
        sock.close()
        raise Exception(f"Прокси отказал в CONNECT туннеле: {status_line}")
        
    return sock

class ProxiedSMTP(smtplib.SMTP):
    def __init__(self, host, port, sock, timeout=None):
        self._my_sock = sock
        super().__init__(timeout=timeout)
        self._host = host
        self.host = host
        self.connect(host, port)
        
    def _get_socket(self, host, port, timeout):
        return self._my_sock

class ProxiedIMAP4_SSL(imaplib.IMAP4_SSL):
    def __init__(self, host, port, ssl_sock, timeout=None):
        self._ssl_sock = ssl_sock
        super().__init__(host, port, timeout=timeout)
        
    def _create_socket(self, timeout):
        return self._ssl_sock

def transliterate_filename(filename):
    cyrillic = 'абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ'
    latin = [
        'a', 'b', 'v', 'g', 'd', 'e', 'yo', 'zh', 'z', 'i', 'y', 'k', 'l', 'm', 'n', 'o', 'p', 'r', 's', 't', 'u', 'f', 'h', 'ts', 'ch', 'sh', 'shch', '', 'y', '', 'e', 'yu', 'ya',
        'A', 'B', 'V', 'G', 'D', 'E', 'Yo', 'Zh', 'Z', 'I', 'Y', 'K', 'L', 'M', 'N', 'O', 'P', 'R', 'S', 'T', 'U', 'F', 'H', 'Ts', 'Ch', 'Sh', 'Shch', '', 'Y', '', 'E', 'Yu', 'Ya'
    ]
    trans_dict = {ord(c): l for c, l in zip(cyrillic, latin)}
    transliterated = filename.translate(trans_dict)
    safe_chars = re.sub(r'[^\x00-\x7F]', '', transliterated)
    return safe_chars.replace(' ', '_')

def send_email_with_attachments(receiver_email, file_paths, vessel_name):
    try:
        msg = EmailMessage()
        msg['Subject'] = f"Ship Documents Pack: {vessel_name} - {datetime.now().strftime('%d.%m.%Y')}"
        msg['From']    = EMAIL_SENDER
        msg['To']      = str(receiver_email).strip()
        msg['Bcc']     = EMAIL_SENDER 
        
        msg.set_content(
            f"Приветствую!\n\nВо вложении направляю пакет документов для судна {vessel_name}.\n\n"
            f"С уважением,\nООО \"ПОДВОДНЫЙ ТЕХНИЧЕСКИЙ ЭКСПЕРТ\""
        )
        final_attachments = list(file_paths)
        has_pdf_report = any(p.lower().endswith('.pdf') for p in file_paths)
        if has_pdf_report and os.path.exists(STATIC_ATTACHMENTS_FOLDER):
            static_files = sorted([
                os.path.join(STATIC_ATTACHMENTS_FOLDER, f)
                for f in os.listdir(STATIC_ATTACHMENTS_FOLDER)
                if os.path.isfile(os.path.join(STATIC_ATTACHMENTS_FOLDER, f))
            ])
            for sf in static_files[:3]:
                if sf not in final_attachments:
                    final_attachments.append(sf)
        for path in final_attachments:
            if not os.path.exists(path):
                continue
            ctype, _ = mimetypes.guess_type(path)
            if ctype is None:
                ctype = 'application/octet-stream'
            maintype, subtype = ctype.split('/', 1)
            
            # Имя файла передаем как ASCII-строку для совместимости с почтовыми клиентами
            safe_filename = transliterate_filename(os.path.basename(path))
            
            with open(path, 'rb') as fp:
                msg.add_attachment(fp.read(), maintype=maintype, subtype=subtype, filename=safe_filename)
                
        if PROXY_URL:
            # На порту 2525 через STARTTLS прокси не блокирует трафик
            plain_sock = connect_plain_socket_via_proxy(SMTP_SERVER, 2525, PROXY_URL, timeout=120)
            server = ProxiedSMTP(SMTP_SERVER, 2525, plain_sock, timeout=120)
            with server:
                server.ehlo()
                server.starttls()
                server.ehlo()
                server.login(EMAIL_SENDER, EMAIL_PASSWORD)
                server.send_message(msg)
        else:
            server = smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT, timeout=120)
            with server:
                server.login(EMAIL_SENDER, EMAIL_PASSWORD)
                server.send_message(msg)
        return True
    except Exception as e:
        log_error(f"Ошибка SMTP при отправке письма на адрес {receiver_email} для судна {vessel_name}", e)
        return False

# =====================================================================
# ПЕРЕСЫЛКА ПИСЬМА ВТОРЫМ БОТОМ (чистый текст без parse_mode уязвимостей)
# =====================================================================
def forward_email_via_second_bot(mail_from, mail_to, mail_cc, mail_subject, mail_body, attachment_paths):
    if not second_bot or not SECOND_BOT_CHAT_ID:
        return

    header_text = (
        f"[ВТОРОЙ БОТ ПОЧТА] Перехвачено входящее письмо:\n\n"
        f"От: {mail_from}\n"
        f"Кому: {mail_to}\n"
        f"Копия: {mail_cc if mail_cc else '—'}\n"
        f"Тема: {mail_subject}\n\n"
        f"Текст письма:\n{mail_body[:3500] if mail_body else '(пусто)'}"
    )

    try:
        second_bot.send_message(SECOND_BOT_CHAT_ID, header_text)
    except Exception as e:
        log_error("Второй бот: ошибка при отправке заголовка письма", e)

    for path in attachment_paths:
        if not os.path.exists(path):
            continue
        try:
            with open(path, 'rb') as f:
                second_bot.send_document(SECOND_BOT_CHAT_ID, f, caption=f"📎 {os.path.basename(path)}")
        except Exception as e:
            log_error(f"Второй бот: ошибка при отправке вложения {path}", e)

# =====================================================================
# МАРКАПЫ БОТА
# =====================================================================
def make_file_selection_markup(files_list, selected, email_dest=None):
    markup = types.InlineKeyboardMarkup(row_width=1)
    for idx, path in enumerate(files_list):
        status_icon = "✅" if selected.get(str(idx), False) else "⬜"
        markup.add(types.InlineKeyboardButton(
            f"{status_icon} {os.path.basename(path)}",
            callback_data=f"togglefile_{idx}"
        ))
    if email_dest:
        markup.add(
            types.InlineKeyboardButton(f"✉️ Отправить на {email_dest}", callback_data="send_to_preselected"),
            types.InlineKeyboardButton("✏️ Изменить email", callback_data="choose_email_dest")
        )
    else:
        markup.add(
            types.InlineKeyboardButton("📧 ВЫБРАТЬ ПОЧТУ ДЛЯ ОТПРАВКИ", callback_data="choose_email_dest")
        )
    markup.add(
        types.InlineKeyboardButton("🗑 ЗАВЕРШИТЬ И ОЧИСТЬ СЕРВЕР",   callback_data="clear_session_data")
    )
    return markup

def make_email_options_markup():
    markup = types.InlineKeyboardMarkup(row_width=1)
    for idx, email_addr in enumerate(EMAIL_OPTIONS):
        markup.add(types.InlineKeyboardButton(f"📬 {email_addr}", callback_data=f"sendtoemail_{idx}"))
    markup.add(types.InlineKeyboardButton("⬅️ Назад к списку файлов", callback_data="back_to_files"))
    return markup

# =====================================================================
# WEB APP — ПОСТРОЕНИЕ URL И ОТПРАВКА КНОПКИ
# =====================================================================
def build_webapp_url(vessel_data):
    """Упаковывает данные судна в base64 и передаёт через параметр ?data="""
    payload = {
        "vessel_name":     vessel_data.get("vessel_name", ""),
        "ship_type":       vessel_data.get("ship_type", ""),
        "imo":             vessel_data.get("imo", ""),
        "flag":            vessel_data.get("flag", ""),
        "owner":           vessel_data.get("owner", ""),
        "master_name":     vessel_data.get("master_name", ""),
        "loa":             str(vessel_data.get("loa", "")),
        "breadth":         str(vessel_data.get("breadth", "")),
        "draft_fore":      str(vessel_data.get("draft_fore", "0.0")),
        "draft_aft":       str(vessel_data.get("draft_aft", "0.0")),
        "act_number":      vessel_data.get("act_number", ""),
        "diving_date":     vessel_data.get("diving_date", ""),
        "date_start":      vessel_data.get("date_start", ""),
        "date_end":        vessel_data.get("date_end", ""),
        "current_date":    vessel_data.get("current_date", datetime.now().strftime("%d.%m.%Y")),
        "wetted_surface":  str(vessel_data.get("wetted_surface", "0.0")),
        "propeller_area":  str(vessel_data.get("propeller_area", "0.0")),
        "company_key":     vessel_data.get("company_key", ""),
        "last_doc_number": load_history_db().get("last_doc_number", 145),
    }
    encoded = base64.b64encode(json.dumps(payload, ensure_ascii=False).encode()).decode()
    # Adding a cache-buster version parameter to force Telegram to load the latest HTML
    return f"{WEBAPP_URL}?v=405&data={encoded}"

def send_webapp_button(chat_id, vessel_data, intro_text=None):
    """Отправляет кнопку Web App как ReplyKeyboardMarkup (для корректной работы sendData) и inline-кнопку текстового режима как фолбэк."""
    v_name = vessel_data.get("vessel_name", "Судно")
    text = intro_text or (
        f"🚢 **{v_name}** — данные распознаны ИИ.\n\n"
        f"Откройте приложение с помощью кнопки **⚡ Открыть редактор судна** на клавиатуре ниже, проверьте поля, выберите компанию и нажмите **Генерировать**."
    )
    
    # Клавиатура (Reply Keyboard) для работы tg.sendData
    reply_markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    reply_markup.add(types.KeyboardButton(
        "⚡ Открыть редактор судна",
        web_app=types.WebAppInfo(url=build_webapp_url(vessel_data))
    ))
    
    # Inline-клавиатура для текстового режима
    inline_markup = types.InlineKeyboardMarkup()
    inline_markup.add(types.InlineKeyboardButton(
        "📝 Текстовый режим (старый)",
        callback_data="textmode"
    ))
    
    bot.send_message(chat_id, text, reply_markup=reply_markup, parse_mode="Markdown")
    bot.send_message(chat_id, "Или используйте текстовый интерфейс:", reply_markup=inline_markup)

def ask_company_choice(chat_id):
    markup = types.InlineKeyboardMarkup(row_width=1)
    for key, c_info in COMPANIES_DATA.items():
        markup.add(types.InlineKeyboardButton(c_info["title"], callback_data=f"selectco_{key}"))
    bot.send_message(chat_id, "🏢 **Выберите фирму для оформления пакета документов:**", reply_markup=markup, parse_mode="Markdown")

def show_review_menu(chat_id):
    d = user_states[chat_id]
    if not d.get('act_number'):
        d['act_number'] = get_next_act_number_for_date(d.get('diving_date', datetime.now().strftime("%d.%m.%Y")))
    text = (
        "📋 **ПРОВЕРКА ДАННЫХ ПЕРЕД ЗАПУСКОМ КОНВЕЙЕРА:**\n\n"
        f"🚢 Судно: `{d.get('vessel_name')}` | ⛴ Тип: `{d.get('ship_type')}`\n"
        f"🔢 IMO: `{d.get('imo')}` | 🏴 Флаг: `{d.get('flag')}`\n"
        f"📏 LOA: `{d.get('loa')} м` / Breadth: `{d.get('breadth')} м`\n"
        f"🌊 Осадка: FWD `{d.get('draft_fore')}` / AFT `{d.get('draft_aft')}`\n"
        f"👤 Капитан: `{d.get('master_name')}`\n"
        f"📜 **Номер Акта осмотра:** `{d.get('act_number')}`\n"
        f"🔢 Номер документов (БД): `{load_history_db().get('last_doc_number', 0) + 1}`\n"
        f"🤿 Число осмотра (подхода): `{d.get('diving_date')}`\n"
        f"📅 Период: `{d.get('date_start')}` - `{d.get('date_end')}`\n\n"
        f"🏢 Фирма: *{d.get('company_short')}*\n\n"
        f"📐 WSA: *{d.get('wetted_surface')}* м² | Винт: *{d.get('propeller_area')}* м²\n"
    )
    markup = types.InlineKeyboardMarkup(row_width=2)
    markup.add(
        types.InlineKeyboardButton("📝 Название судна",    callback_data="edit_vessel_name"),
        types.InlineKeyboardButton("📝 Осадки FWD/AFT",   callback_data="edit_draft"),
        types.InlineKeyboardButton("📝 Капитан ФИО",      callback_data="edit_master_name"),
        types.InlineKeyboardButton("📝 Акт / Даты",       callback_data="edit_act_dates"),
        types.InlineKeyboardButton("📝 Параметры / Флаг", callback_data="edit_vessel_specs"),
        types.InlineKeyboardButton("🏢 Сменить Фирму",    callback_data="edit_company_choice"),
        types.InlineKeyboardButton("✅ ВСЁ ВЕРНО - СГЕНЕРИРОВАТЬ ПАКЕТ", callback_data="ask_start_num")
    )
    bot.send_message(chat_id, text, reply_markup=markup, parse_mode="Markdown")

def generate_all_documents(chat_id, start_num, msg_to_edit, task_id=None):
    try:
        vessel_data   = user_states[chat_id]
        vessel_name   = re.sub(r'[\\/*?:"<>|]', "_", str(vessel_data.get('vessel_name', 'UNKNOWN')).strip())
        vessel_output_folder = os.path.join(BASE_DIR, "READY_DOCUMENTS", vessel_name)
        os.makedirs(vessel_output_folder, exist_ok=True)
        template_files = sorted([f for f in os.listdir(TEMPLATE_FOLDER) if f.lower().endswith(".docx")])
        try:
            current_doc_num = int(start_num)
        except:
            current_doc_num = 145

        if msg_to_edit is None:
            try:
                msg_to_edit = bot.send_message(chat_id, "⚙️ Запуск генерации...")
            except:
                pass

        generated_paths = []
        docx_to_convert = []
        for file_idx, template_file in enumerate(template_files, start=1):
            if task_id:
                update_task_status(
                    task_id, 
                    "generating_docs", 
                    10 + int(25 * (file_idx / len(template_files))), 
                    f"Генерация {template_file}..."
                )
            try:
                if msg_to_edit:
                    bot.edit_message_text(
                        f"📝 Генерирую документ ({file_idx}/{len(template_files)}):\n`{template_file}`...",
                        chat_id, msg_to_edit.message_id, parse_mode="Markdown"
                    )
            except:
                pass
            doc = Document(os.path.join(TEMPLATE_FOLDER, template_file))
            file_has_doc_num_tag = (
                any("{{doc_number}}" in p.text for p in doc.paragraphs) or
                any(
                    any("{{doc_number}}" in cell.text for cell in row.cells)
                    for table in doc.tables for row in table.rows
                )
            )
            try:
                raw_aft = str(vessel_data.get("draft_aft", "0.0")).replace(',', '.').strip()
                numbers = re.findall(r"\d+\.\d+|\d+", raw_aft)
                da = float(numbers[0]) if numbers else 0.0
                draft_aft_round = int(da) + 1 if da % 1 == 0 else math.ceil(da)
            except:
                draft_aft_round = "8"
            full_diving_date = str(vessel_data.get("diving_date", "")).strip()
            full_row_dict = {
                "vessel_name": str(vessel_data.get("vessel_name", "")),
                "ship_type":   str(vessel_data.get("ship_type", "")),
                "imo":         str(vessel_data.get("imo", "")),
                "loa":         str(vessel_data.get("loa", "")),
                "breadth":     str(vessel_data.get("breadth", "")),
                "flag":        str(vessel_data.get("flag", "")),
                "owner":       str(vessel_data.get("owner", "")),
                "draft_fore":  str(vessel_data.get("draft_fore", "")),
                "draft_aft":   str(vessel_data.get("draft_aft", "")),
                "dra": str(draft_aft_round), "Dra": str(draft_aft_round), "DRA": str(draft_aft_round),
                "master_name":       str(vessel_data.get("master_name", "")),
                "date_start":        str(vessel_data.get("date_start", "")),
                "date_end":          str(vessel_data.get("date_end", "")),
                "act_number":        str(vessel_data.get("act_number", "")),
                "current_date":      str(vessel_data.get("current_date", "")),
                "diving_date":       full_diving_date,
                "wsa":               str(vessel_data.get("wetted_surface", 0.0)),
                "propeller_area":    str(vessel_data.get("propeller_area", 0.0)),
                "company_contract":  str(vessel_data.get("company_contract", "")),
                "company_short":     str(vessel_data.get("company_short", "")),
                "company_full":      str(vessel_data.get("company_full", "")),
                "doc_number":        str(current_doc_num) if file_has_doc_num_tag else ""
            }
            global_tag_replacer_mono(doc, full_row_dict)
            docx_output_path = os.path.join(vessel_output_folder, template_file)
            doc.save(docx_output_path)
            generated_paths.append(docx_output_path)
            if file_has_doc_num_tag:
                current_doc_num += 1
            if file_idx >= PDF_START_FROM_FILE:
                pdf_path = docx_output_path.replace(".docx", ".pdf")
                docx_to_convert.append((docx_output_path, pdf_path))

        # Выполняем пакетную конвертацию всех DOCX в PDF
        if docx_to_convert:
            if task_id:
                update_task_status(
                    task_id, 
                    "converting_pdfs", 
                    45, 
                    f"Конвертация {len(docx_to_convert)} файлов в PDF..."
                )
            if platform.system() == "Windows":
                for docx_p, pdf_p in docx_to_convert:
                    run_isolated_conversion(docx_p, pdf_p)
            else:
                # Linux / Render: Запускаем LibreOffice ОДИН раз для всех файлов, ускоряя процесс в 3-4 раза!
                import subprocess
                try:
                    outdir = os.path.dirname(docx_to_convert[0][1])
                    docx_list = [item[0] for item in docx_to_convert]
                    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir] + docx_list
                    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
                    if result.returncode != 0:
                        log_error(f"Сбой пакетной конвертации LibreOffice ({result.returncode}): {result.stderr.decode('utf-8', errors='ignore')}", None)
                except Exception as e:
                    log_error("Сбой пакетной конвертации через LibreOffice", e)
            
            # Добавляем успешно созданные PDF в список сгенерированных
            for docx_p, pdf_p in docx_to_convert:
                if os.path.exists(pdf_p):
                    generated_paths.append(pdf_p)
        try:
            if msg_to_edit:
                bot.delete_message(chat_id, msg_to_edit.message_id)
        except:
            pass
        archive_folder = os.path.join(BASE_DIR, "FINAL_ARCHIVE", f"{vessel_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        os.makedirs(archive_folder, exist_ok=True)
        for p in generated_paths:
            if os.path.exists(p):
                shutil.copy(p, archive_folder)
        
        if task_id:
            update_task_status(task_id, "sending_zip", 75, "Создание архива...")
            
        zip_filename = f"DOCUMENTS_{vessel_name}_{task_id}" if task_id else f"DOCUMENTS_{vessel_name}"
        zip_base = os.path.join(BASE_DIR, "READY_DOCUMENTS", zip_filename)
        shutil.make_archive(zip_base, 'zip', vessel_output_folder)
        
        # Send to telegram chat as backup/log
        try:
            with open(f"{zip_base}.zip", 'rb') as z_file:
                bot.send_document(chat_id, z_file, caption=f"✅ Пакет документов для {vessel_name} готов!", timeout=120)
        except Exception as tg_err:
            log_error("Не удалось отправить ZIP в Telegram (продолжаем работу):", tg_err)

        if task_id:
            # Save files list and ZIP path in task status so the Web App can query and download them
            generation_tasks[task_id]["zip_path"] = f"{zip_base}.zip"
            generation_tasks[task_id]["files"] = [os.path.basename(p) for p in generated_paths]
            generation_tasks[task_id]["vessel_folder"] = vessel_output_folder
            generation_tasks[task_id]["vessel_name"] = vessel_name
        else:
            try:
                os.remove(f"{zip_base}.zip")
            except:
                pass
        log_inspection_event(current_doc_num - 1, vessel_data.get('act_number', ''), vessel_name, vessel_data.get('diving_date', ''), vessel_data)
        default_selected = {str(i): path.lower().endswith('.pdf') for i, path in enumerate(generated_paths)}
        user_states[chat_id].update({
            "generated_files": list(generated_paths),
            "vessel_folder":   vessel_output_folder,
            "selected_files":  default_selected,
            "step":            "waiting_to_send"
        })
        
        if task_id:
            update_task_status(
                task_id, 
                "waiting_for_confirmation", 
                85, 
                "Архив готов и отправлен в Telegram. Ожидаю подтверждения отправки почты."
            )

        email_dest = vessel_data.get("email_dest", "")
        markup = types.InlineKeyboardMarkup(row_width=1)
        if email_dest:
            markup.add(
                types.InlineKeyboardButton(f"🚀 Отправить всё на {email_dest}", callback_data="send_all_to_preset"),
                types.InlineKeyboardButton("⚙️ Настроить вручную", callback_data="start_email_delivery")
            )
            msg_text = f"📂 Пакет документов готов и отправлен выше в архиве.\n\nПожалуйста, ознакомьтесь с ним. Вы можете отправить все документы на выбранный адрес **{email_dest}** или выбрать файлы вручную."
        else:
            markup.add(types.InlineKeyboardButton("✉️ Перейти к отправке на почту", callback_data="start_email_delivery"))
            msg_text = "📂 Пакет документов готов и отправлен выше в архиве.\n\nПожалуйста, ознакомьтесь с ним. Если всё верно, нажмите кнопку ниже для выбора файлов и отправки на почту."

        bot.send_message(
            chat_id,
            msg_text,
            reply_markup=markup,
            parse_mode="Markdown"
        )
    except Exception as e:
        log_error("Критическая поломка во время генерации пакета Word/PDF", e)
        if task_id:
            update_task_status(task_id, "failed", 100, f"Критическая ошибка генерации: {str(e)}", error=str(e))
        bot.send_message(chat_id, "❌ Произошла системная ошибка при сборке пакета документов. Ошибка записана в логгер.")

# =====================================================================
# ИНТЕЛЛЕКТУАЛЬНЫЙ ПАРСЕР ИИ
# =====================================================================
def analyze_combined_email_data(text_body, pdf_paths):
    base64_images = []
    for path in pdf_paths:
        try:
            base64_images.extend(extract_base64_images_from_pdf(path))
        except:
            pass
    prompt_text = f"""
    You are a robotic marine data intelligence extraction API. Form a valid JSON object by parsing the context.
    EMAIL RAW TEXT:
    {text_body}
    CRITICAL INSTRUCTIONS:
    1. 'vessel_name': Pure ship title (UPPERCASE, remove MV/MT/MV.).
    2. 'ship_type': Exactly one of: "TANKER", "BULK CARRIER", "GENERAL CARGO", "CONTAINER".
    3. 'loa': Overall length of vessel (float, e.g. 141.2).
    4. 'breadth': Extreme breadth of vessel (float, e.g. 22.8).
    5. 'imo': IMO number (7 digits).
    6. 'flag': Pure registration country name (e.g. "MALTA", "PANAMA", "TURKEY"). MUST EXTRACT. DO NOT OMIT.
    7. 'owner': Registered shipowner company name.
    8. 'draft_fore' & 'draft_aft': Floats. If missing, set 0.0.
    9. 'diving_date': DD.MM.YYYY. If missing, set to current date.
    10. 'master_name': UPPERCASE formatted "GIVEN_NAME FAMILY_NAME".
    """
    content_array = [{"type": "text", "text": prompt_text}]
    for b64 in base64_images:
        content_array.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
    try:
        res = client.chat.completions.create(
            model=CURRENT_MODEL,
            messages=[{"role": "user", "content": content_array}],
            temperature=0.0,
            timeout=60.0
        )
        return json.loads(extract_json_from_text(res.choices[0].message.content.strip()))
    except Exception as e:
        log_error("Ошибка ИИ-парсера при попытке распознать спецификацию", e)
        return None

# =====================================================================
# ФОНОВЫЙ ПОТОК (ДЕМОН) СБОРА ПОЧТЫ — С BACKOFF И ПЕРЕСЫЛКОЙ
# =====================================================================
def mail_checker_daemon():
    backoff_delays = [30, 60, 120, 300]  
    fail_count = 0

    while True:
        if not BOT_ACTIVE:
            time.sleep(5)
            continue
        try:
            if PROXY_URL:
                ssl_sock = connect_ssl_via_proxy(IMAP_SERVER, IMAP_PORT, PROXY_URL, timeout=120)
                mail = ProxiedIMAP4_SSL(IMAP_SERVER, IMAP_PORT, ssl_sock, timeout=120)
            else:
                mail = imaplib.IMAP4_SSL(IMAP_SERVER, IMAP_PORT)
            mail.login(EMAIL_SENDER, EMAIL_PASSWORD)
            mail.select("inbox")
            status, response = mail.search(None, 'UNSEEN')
            if status != 'OK':
                mail.logout()
                fail_count = 0
                time.sleep(30)
                continue

            for e_id in response[0].split():
                mail.store(e_id, '+FLAGS', '\\Seen')

                _, data = mail.fetch(e_id, '(RFC822)')
                msg = email.message_from_bytes(data[0][1])

                def decode_hdr(raw):
                    if not raw:
                        return ""
                    parts = decode_header(raw)
                    result = ""
                    for part, enc in parts:
                        if isinstance(part, bytes):
                            result += part.decode(enc or "utf-8", errors="ignore")
                        else:
                            result += str(part)
                    return result

                mail_subject = decode_hdr(msg.get("Subject", ""))
                mail_from    = decode_hdr(msg.get("From", ""))
                mail_to      = decode_hdr(msg.get("To", ""))
                mail_cc      = decode_hdr(msg.get("Cc", ""))

                body_text   = ""
                temp_pdf_paths  = []   
                all_attach_paths = []  

                if msg.is_multipart():
                    for part in msg.walk():
                        content_disp = str(part.get("Content-Disposition", ""))
                        name_raw     = part.get_filename()

                        if "attachment" in content_disp and name_raw:
                            h_name, e_enc = decode_header(name_raw)[0]
                            if isinstance(h_name, bytes):
                                h_name = h_name.decode(e_enc or "utf-8", errors="ignore")

                            t_path = os.path.join(BASE_DIR, f"mail_temp_{h_name}")
                            with open(t_path, "wb") as f:
                                f.write(part.get_payload(decode=True))

                            if h_name.lower().endswith(".pdf"):
                                temp_pdf_paths.append(t_path)

                            if not is_image_file(h_name):
                                all_attach_paths.append(t_path)

                        elif part.get_content_type() == "text/plain":
                            try:
                                body_text += part.get_payload(decode=True).decode("utf-8", errors="ignore")
                            except:
                                pass
                else:
                    body_text = msg.get_payload(decode=True).decode("utf-8", errors="ignore")

                forward_email_via_second_bot(
                    mail_from, mail_to, mail_cc,
                    mail_subject, body_text,
                    all_attach_paths
                )

                data_ai = analyze_combined_email_data(body_text, temp_pdf_paths)

                for p in temp_pdf_paths:
                    if os.path.exists(p):
                        os.remove(p)
                for p in all_attach_paths:
                    if os.path.exists(p) and p not in temp_pdf_paths:
                        os.remove(p)

                if not data_ai:
                    continue

                v_name_raw = str(data_ai.get("vessel_name", "UNKNOWN")).upper()
                div_date   = data_ai.get("diving_date", datetime.now().strftime("%d.%m.%Y"))
                if "DD.MM.YYYY" in div_date or not div_date:
                    div_date = datetime.now().strftime("%d.%m.%Y")

                try:
                    dt_obj         = datetime.strptime(div_date.strip(), "%d.%m.%Y")
                    start_dt       = dt_obj - timedelta(days=1)
                    date_start_str = start_dt.strftime("%d.%m.%Y")
                    date_end_str   = (start_dt + timedelta(days=4)).strftime("%d.%m.%Y")
                except:
                    date_start_str = datetime.now().strftime("%d.%m.%Y")
                    date_end_str   = (datetime.now() + timedelta(days=4)).strftime("%d.%m.%Y")

                session_key = f"mail_{v_name_raw}_{int(time.time())}"
                user_states[session_key] = {
                    "vessel_name":  v_name_raw,
                    "ship_type":    str(data_ai.get("ship_type", "GENERAL CARGO")).upper(),
                    "imo":          str(data_ai.get("imo", "0000000")),
                    "loa":          str(data_ai.get("loa", "0.0")),
                    "breadth":      str(data_ai.get("breadth", "0.0")),
                    "flag":         str(data_ai.get("flag", "НЕ УКАЗАН")).upper(),
                    "owner":        str(data_ai.get("owner", "НЕ УКАЗАН")).upper(),
                    "draft_fore":   str(data_ai.get("draft_fore", "0.0")),
                    "draft_aft":    str(data_ai.get("draft_aft", "0.0")),
                    "master_name":  str(data_ai.get("master_name", "НЕ НАЙДЕН")).upper(),
                    "act_number":   get_next_act_number_for_date(div_date),
                    "diving_date":  div_date,
                    "date_start":   date_start_str,
                    "date_end":     date_end_str,
                    "current_date": datetime.now().strftime("%d.%m.%Y"),
                    "company_contract": "", "company_short": "НЕ ВЫБРАНА", "company_full": ""
                }
                update_draft_calculations(
                    user_states[session_key],
                    user_states[session_key]["draft_fore"],
                    user_states[session_key]["draft_aft"]
                )

                reply_markup = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
                reply_markup.add(types.KeyboardButton(
                    "⚡ Открыть редактор судна",
                    web_app=types.WebAppInfo(url=build_webapp_url(user_states[session_key]))
                ))

                inline_markup = types.InlineKeyboardMarkup()
                inline_markup.add(types.InlineKeyboardButton(
                    "📥 Взять в работу (текстовый режим)",
                    callback_data=f"loadmail_{session_key}"
                ))

                bot.send_message(
                    MY_CHAT_ID,
                    f"📥 **На почту пришло судно: {v_name_raw}**\n\n"
                    f"Используйте кнопку **⚡ Открыть редактор судна** на клавиатуре ниже для запуска редактора.",
                    reply_markup=reply_markup,
                    parse_mode="Markdown"
                )
                bot.send_message(
                    MY_CHAT_ID,
                    "Или нажмите для текстового режима:",
                    reply_markup=inline_markup
                )

            mail.close()
            mail.logout()

            fail_count = 0
            time.sleep(30)

        except Exception as e:
            log_error("Критический сбой внутри циклического фонового IMAP демона почты", e)
            delay = backoff_delays[min(fail_count, len(backoff_delays) - 1)]
            fail_count += 1
            time.sleep(delay)

# =====================================================================
# ОБРАБОТЧИК WEB APP DATA — СЕРДЦЕ ИНТЕГРАЦИИ
# =====================================================================
@bot.message_handler(content_types=['web_app_data'])
def handle_webapp_data(message):
    """Получаем данные из Web App после нажатия кнопки 'Генерировать'."""
    chat_id = message.chat.id
    try:
        payload = json.loads(message.web_app_data.data)
    except Exception as e:
        log_error("WebApp: ошибка парсинга данных", e)
        bot.send_message(chat_id, "❌ Ошибка при получении данных из Web App.")
        return

    action = payload.get("action", "generate")
    if action == "cancel":
        bot.send_message(chat_id, "↩️ Работа с судном отменена.")
        return
    if action != "generate":
        bot.send_message(chat_id, "⚠️ Неизвестная команда из Web App.")
        return

    company_key = payload.get("company_key", "")
    co = COMPANIES_DATA.get(company_key, {})

    vessel_data = {
        "vessel_name":      str(payload.get("vessel_name", "")).upper(),
        "ship_type":        str(payload.get("ship_type", "GENERAL CARGO")).upper(),
        "imo":              str(payload.get("imo", "")),
        "flag":             str(payload.get("flag", "")).upper(),
        "owner":            str(payload.get("owner", "")).upper(),
        "master_name":      str(payload.get("master_name", "")).upper(),
        "loa":              str(payload.get("loa", "")),
        "breadth":          str(payload.get("breadth", "")),
        "draft_fore":       str(payload.get("draft_fore", "0.0")),
        "draft_aft":        str(payload.get("draft_aft", "0.0")),
        "act_number":       str(payload.get("act_number", "")),
        "diving_date":      str(payload.get("diving_date", "")),
        "date_start":       str(payload.get("date_start", "")),
        "date_end":         str(payload.get("date_end", "")),
        "current_date":     str(payload.get("current_date", datetime.now().strftime("%d.%m.%Y"))),
        "wetted_surface":   str(payload.get("wetted_surface", "0.0")),
        "propeller_area":   str(payload.get("propeller_area", "0.0")),
        "company_contract": co.get("company_contract", ""),
        "company_short":    co.get("company_short", "НЕ ВЫБРАНА"),
        "company_full":     co.get("company_full", ""),
        "email_dest":       str(payload.get("email_dest", "")),
    }

    start_num = int(payload.get("start_num", load_history_db().get("last_doc_number", 145) + 1))
    user_states[chat_id] = vessel_data

    status_msg = bot.send_message(chat_id, "⚙️ Запускаю генерацию документов из Web App...", reply_markup=types.ReplyKeyboardRemove())
    generate_all_documents(chat_id, start_num, status_msg)

# =====================================================================
# ОБРАБОТЧИКИ ТЕЛЕГРАМ-КОМАНД
# =====================================================================
@bot.message_handler(commands=['start', 'status', 'acts'])
def handle_commands(message):
    try:
        if message.text.startswith('/status'):
            bot.send_message(message.chat.id, compile_advanced_status_text(), parse_mode="Markdown")
        elif message.text.startswith('/acts'):
            if not os.path.exists(COUNTERS_FILE):
                bot.send_message(message.chat.id, "📭 **БД отсутствует!**", parse_mode="Markdown")
                return
            with open(COUNTERS_FILE, 'r', encoding='utf-8') as f:
                db = json.load(f)
            acts_list = db.get("used_acts", [])
            if not acts_list:
                bot.send_message(message.chat.id, "📭 **История пуста!**", parse_mode="Markdown")
                return
            text = f"📋 **ПОСЛЕДНИЕ АКТЫ В БАЗЕ (Всего: {len(acts_list)}):**\n\n"
            for idx, a in enumerate(acts_list[-10:][::-1], start=1):
                text += f"{idx}. 📜 Акт: `{a.get('act_number')}` | 🚢 Судно: *{str(a.get('vessel_name')).upper()}* | 🤿 Дата: `{a.get('diving_date')}`\n"
            bot.send_message(message.chat.id, text, parse_mode="Markdown")
        else:
            bot.send_message(
                message.chat.id,
                "⚓️ **Конвейер V3.9 готов!** Отправьте PDF спецификации судна вручную или ждите писем.",
                parse_mode="Markdown"
            )
    except Exception as e:
        log_error("Ошибка при обработке текстовой команды TG бота", e)

@bot.message_handler(content_types=['document'])
def handle_manual_document(message):
    chat_id = message.chat.id
    try:
        if not message.document.file_name.lower().endswith('.pdf'):
            bot.send_message(chat_id, "⚠️ Только файлы PDF!")
            return
        file_info = bot.get_file(message.document.file_id)
        t_path = os.path.join(BASE_DIR, f"manual_{message.document.file_name}")
        with open(t_path, "wb") as f:
            f.write(bot.download_file(file_info.file_path))
        msg_status = bot.send_message(chat_id, "🔍 **Ручной запуск!** Анализирую спецификацию через ИИ...", parse_mode="Markdown")
        data_ai = analyze_combined_email_data("", [t_path])
        if os.path.exists(t_path):
            os.remove(t_path)
        if not data_ai:
            return
        div_date = data_ai.get("diving_date", datetime.now().strftime("%d.%m.%Y"))
        if "DD.MM.YYYY" in div_date or not div_date:
            div_date = datetime.now().strftime("%d.%m.%Y")
        try:
            dt_obj         = datetime.strptime(div_date.strip(), "%d.%m.%Y")
            date_start_str = (dt_obj - timedelta(days=1)).strftime("%d.%m.%Y")
            date_end_str   = (dt_obj + timedelta(days=3)).strftime("%d.%m.%Y")
        except:
            date_start_str = datetime.now().strftime("%d.%m.%Y")
            date_end_str   = (datetime.now() + timedelta(days=4)).strftime("%d.%m.%Y")
        user_states[chat_id] = {
            "vessel_name":  str(data_ai.get("vessel_name", "UNKNOWN")).upper(),
            "ship_type":    str(data_ai.get("ship_type", "GENERAL CARGO")).upper(),
            "imo":          str(data_ai.get("imo", "0000000")),
            "loa":          str(data_ai.get("loa", "0.0")),
            "breadth":      str(data_ai.get("breadth", "0.0")),
            "flag":         str(data_ai.get("flag", "НЕ УКАЗАН")).upper(),
            "owner":        str(data_ai.get("owner", "НЕ УКАЗАН")).upper(),
            "draft_fore":   str(data_ai.get("draft_fore", "0.0")),
            "draft_aft":    str(data_ai.get("draft_aft", "0.0")),
            "master_name":  str(data_ai.get("master_name", "НЕ НАЙДЕН")).upper(),
            "act_number":   get_next_act_number_for_date(div_date),
            "diving_date":  div_date,
            "date_start":   date_start_str,
            "date_end":     date_end_str,
            "current_date": datetime.now().strftime("%d.%m.%Y"),
            "company_contract": "", "company_short": "НЕ ВЫБРАНА", "company_full": ""
        }
        update_draft_calculations(user_states[chat_id], user_states[chat_id]["draft_fore"], user_states[chat_id]["draft_aft"])
        try:
            bot.delete_message(chat_id, msg_status.message_id)
        except:
            pass
        send_webapp_button(chat_id, user_states[chat_id])
    except Exception as e:
        log_error("Ошибка при ручной загрузке PDF документа пользователем", e)

# =====================================================================
# ОБРАБОТКА CALLBACK КНОПОК
# =====================================================================
@bot.callback_query_handler(func=lambda call: True)
def handle_callbacks(call):
    chat_id = call.message.chat.id
    data    = call.data
    bot.answer_callback_query(call.id)

    if data.startswith("loadmail_"):
        try:
            session_key = data.replace("loadmail_", "")
            if session_key in user_states:
                user_states[chat_id] = dict(user_states[session_key])
                del user_states[session_key]
                try:
                    bot.delete_message(chat_id, call.message.message_id)
                except:
                    pass
                ask_company_choice(chat_id)
            else:
                bot.send_message(chat_id, "⚠️ Данные этой сессии устарели или были удалены.")
        except Exception as e:
            log_error("Callback loadmail: ошибка при загрузке почтовой сессии", e)
            bot.send_message(chat_id, "❌ Ошибка при загрузке сессии.")
        return

    if chat_id not in user_states:
        bot.send_message(
            chat_id, 
            "⚠️ **Сессия была сброшена из-за длительного бездействия или перезагрузки сервера.**\nПожалуйста, начните работу с пакетом заново (отправьте PDF или выберите из почты повторно).", 
            parse_mode="Markdown"
        )
        return

    if data.startswith("selectco_"):
        try:
            co = COMPANIES_DATA.get(data.replace("selectco_", ""))
            if co:
                user_states[chat_id].update({
                    "company_contract": co["company_contract"],
                    "company_short":    co["company_short"],
                    "company_full":     co["company_full"]
                })
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
            show_review_menu(chat_id)
        except Exception as e:
            log_error("Callback selectco: ошибка выбора компании", e)
            bot.send_message(chat_id, "❌ Ошибка при выборе компании.")

    elif data == "edit_company_choice":
        try:
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
            ask_company_choice(chat_id)
        except Exception as e:
            log_error("Callback edit_company_choice: ошибка", e)

    elif data == "ask_start_num":
        try:
            user_states[chat_id]['step'] = 'waiting_for_start_num'
            last_num = load_history_db().get("last_doc_number", 0)
            msg = bot.send_message(
                chat_id,
                f"🔢 Введите **НАЧАЛЬНЫЙ НОМЕР** документов.\n Рекомендуется: `{last_num + 1}`",
                parse_mode="Markdown"
            )
            user_states[chat_id]['msg_to_edit'] = msg
        except Exception as e:
            log_error("Callback ask_start_num: ошибка", e)

    elif data.startswith("edit_"):
        try:
            target = data.replace("edit_", "")
            user_states[chat_id]['step'] = f"editing_{target}"
            prompts = {
                "vessel_name":  "✍️ Введите новое НАЗВАНИЕ СУДНА:",
                "draft":        "✍️ Введите новые ОСАДКИ (FWD и AFT через пробел):",
                "master_name":  "✍️ Введите ФИО КАПИТАНА:",
                "act_dates":    "✍️ Введите полный формат (Номер акта, Число осмотра ДД.ММ.ГГГГ, Дата старта ДД.ММ.ГГГГ через запятую):",
                "vessel_specs": "✍️ Введите технические параметры через запятую:\n*(IMO, LOA, Breadth, Flag, Owner, Тип судна)*"
            }
            bot.send_message(chat_id, prompts.get(target, "Введите значение:"))
        except Exception as e:
            log_error(f"Callback edit_{data}: ошибка", e)

    elif data.startswith("togglefile_"):
        try:
            idx_str = data.replace("togglefile_", "")
            sel = user_states[chat_id].get("selected_files", {})
            if idx_str in sel:
                sel[idx_str] = not sel[idx_str]
            bot.edit_message_reply_markup(
                chat_id, call.message.message_id,
                reply_markup=make_file_selection_markup(
                    user_states[chat_id]["generated_files"],
                    sel,
                    user_states[chat_id].get("email_dest")
                )
            )
        except Exception as e:
            log_error("Callback togglefile: ошибка", e)

    elif data == "choose_email_dest":
        try:
            bot.send_message(
                chat_id,
                "🏢 **Выберите адрес электронной почты для отправки:**",
                reply_markup=make_email_options_markup(),
                parse_mode="Markdown"
            )
        except Exception as e:
            log_error("Callback choose_email_dest: ошибка", e)

    elif data == "back_to_files":
        try:
            bot.send_message(
                chat_id,
                "📬 **Выбор файлов перед отправкой агенту:**",
                reply_markup=make_file_selection_markup(
                    user_states[chat_id]["generated_files"],
                    user_states[chat_id]["selected_files"],
                    user_states[chat_id].get("email_dest")
                ),
                parse_mode="Markdown"
            )
        except Exception as e:
            log_error("Callback back_to_files: ошибка", e)

    elif data.startswith("sendtoemail_"):
        try:
            chosen_email = EMAIL_OPTIONS[int(data.replace("sendtoemail_", ""))]
            d = user_states[chat_id]
            p_to_send = [d["generated_files"][int(i)] for i, val in d["selected_files"].items() if val]
            if not p_to_send:
                bot.send_message(chat_id, "⚠️ Не выбран ни один файл для отправки.")
                return
            bot.edit_message_text(f"📡 Отправляю файлы на `{chosen_email}`...", chat_id, call.message.message_id, parse_mode="Markdown")
            success = send_email_with_attachments(chosen_email, p_to_send, d.get("vessel_name", "Vessel"))
            bot.send_message(
                chat_id,
                f"✅ **Пакет отправлен на:** `{chosen_email}`" if success else "❌ Ошибка SMTP сервера.",
                parse_mode="Markdown"
            )
            next_selected = {str(i): (i < 3) for i in range(len(d["generated_files"]))}
            d["selected_files"] = next_selected
            bot.send_message(
                chat_id,
                "📬 **Отправка завершена.** Галочки перевыставлены на DOCX документы:",
                reply_markup=make_file_selection_markup(d["generated_files"], next_selected, d.get("email_dest")),
                parse_mode="Markdown"
            )
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
        except Exception as e:
            log_error("Callback sendtoemail: ошибка при отправке на email", e)
            bot.send_message(chat_id, "❌ Ошибка при отправке письма.")

    elif data == "clear_session_data":
        try:
            if folder := user_states[chat_id].get("vessel_folder"):
                try:
                    shutil.rmtree(folder)
                except:
                    pass
            del user_states[chat_id]
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
            bot.send_message(chat_id, "🔒 Сессия закрыта. Данные стерты!", parse_mode="Markdown")
        except Exception as e:
            log_error("Callback clear_session_data: ошибка при очистке", e)

    elif data == "textmode":
        try:
            ask_company_choice(chat_id)
        except Exception as e:
            log_error("Callback textmode: ошибка перехода в текстовый режим", e)

    elif data == "start_email_delivery":
        try:
            user_states[chat_id]["step"] = "file_selection"
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
            bot.send_message(
                chat_id,
                "📬 **Выбор файлов для отправки:**",
                reply_markup=make_file_selection_markup(
                    user_states[chat_id]["generated_files"],
                    user_states[chat_id]["selected_files"],
                    user_states[chat_id].get("email_dest")
                ),
                parse_mode="Markdown"
            )
        except Exception as e:
            log_error("Callback start_email_delivery: ошибка", e)

    elif data == "send_to_preselected":
        try:
            d = user_states[chat_id]
            chosen_email = d.get("email_dest")
            if not chosen_email:
                bot.send_message(chat_id, "⚠️ Email не выбран.")
                return
            p_to_send = [d["generated_files"][int(i)] for i, val in d["selected_files"].items() if val]
            if not p_to_send:
                bot.send_message(chat_id, "⚠️ Не выбран ни один файл для отправки.")
                return
            bot.edit_message_text(f"📡 Отправляю файлы на `{chosen_email}`...", chat_id, call.message.message_id, parse_mode="Markdown")
            success = send_email_with_attachments(chosen_email, p_to_send, d.get("vessel_name", "Vessel"))
            bot.send_message(
                chat_id,
                f"✅ **Пакет отправлен на:** `{chosen_email}`" if success else "❌ Ошибка SMTP сервера.",
                parse_mode="Markdown"
            )
            next_selected = {str(i): (i < 3) for i in range(len(d["generated_files"]))}
            d["selected_files"] = next_selected
            bot.send_message(
                chat_id,
                "📬 **Отправка завершена.** Галочки перевыставлены на DOCX документы:",
                reply_markup=make_file_selection_markup(d["generated_files"], next_selected, d.get("email_dest")),
                parse_mode="Markdown"
            )
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
        except Exception as e:
            log_error("Callback send_to_preselected: ошибка", e)
            bot.send_message(chat_id, "❌ Ошибка при отправке письма.")

    elif data == "send_all_to_preset":
        try:
            d = user_states[chat_id]
            email_dest = d.get("email_dest")
            if not email_dest:
                bot.send_message(chat_id, "⚠️ Адрес электронной почты не задан в сессии.")
                return
            
            generated = d.get("generated_files", [])
            if not generated:
                bot.send_message(chat_id, "⚠️ Нет сгенерированных файлов.")
                return
            
            bot.edit_message_text(f"📡 Отправка документов на `{email_dest}`...", chat_id, call.message.message_id, parse_mode="Markdown")
            
            # 1. Первый пакет (PDF документы)
            pdfs = [p for p in generated if p.lower().endswith('.pdf')]
            files_to_send1 = pdfs if pdfs else generated[:3]
            success1 = send_email_with_attachments(email_dest, files_to_send1, d.get("vessel_name", "Vessel"))
            
            # 2. Второй пакет (первые 3 DOCX документа - Акты)
            docx_files = generated[:3]
            success2 = False
            if docx_files:
                success2 = send_email_with_attachments(email_dest, docx_files, d.get("vessel_name", "Vessel"))
            
            if success1 and success2:
                bot.send_message(chat_id, f"✅ Оба пакета успешно отправлены на `{email_dest}`!", parse_mode="Markdown")
            else:
                msg = "⚠️ Возникли проблемы при отправке:\n"
                msg += f"Пакет 1 (PDF): {'Отправлен' if success1 else 'Ошибка'}\n"
                msg += f"Пакет 2 (DOCX): {'Отправлен' if success2 else 'Ошибка'}"
                bot.send_message(chat_id, msg, parse_mode="Markdown")
                
            try:
                bot.delete_message(chat_id, call.message.message_id)
            except:
                pass
        except Exception as e:
            log_error("Callback send_all_to_preset: ошибка", e)
            bot.send_message(chat_id, "❌ Произошла ошибка при отправке.")

# =====================================================================
# ОБРАБОТЧИКИ ТЕКСТОВЫХ ПРАВОК
# =====================================================================
@bot.message_handler(func=lambda msg: str(user_states.get(msg.chat.id, {}).get('step')).startswith('editing_'))
def handle_text_edits(message):
    chat_id = message.chat.id
    try:
        vessel_data = user_states[chat_id]
        target = vessel_data['step'].replace("editing_", "")
        text   = message.text.strip()
        old_act = vessel_data.get('act_number')

        if target == "vessel_name":
            vessel_data['vessel_name'] = text.upper()
        elif target == "master_name":
            vessel_data['master_name'] = text.upper()
        elif target == "draft":
            try:
                p = text.replace(',', '.').split()
                update_draft_calculations(vessel_data, p[0], p[1])
            except:
                pass
        elif target == "act_dates":
            p = [i.strip() for i in text.replace(',', '\n').split('\n') if i.strip()]
            if len(p) >= 3:
                vessel_data.update({'act_number': p[0], 'diving_date': p[1], 'date_start': p[2]})
                try:
                    dt_obj = datetime.strptime(p[2].strip(), "%d.%m.%Y")
                    vessel_data['date_end'] = (dt_obj + timedelta(days=4)).strftime("%d.%m.%Y")
                except:
                    vessel_data['date_end'] = p[2]
            update_draft_calculations(vessel_data, vessel_data.get("draft_fore", "0.0"), vessel_data.get("draft_aft", "0.0"))
        elif target == "vessel_specs":
            p = [i.strip() for i in text.split(',')]
            if len(p) >= 6:
                vessel_data.update({
                    "imo": p[0], "loa": p[1], "breadth": p[2],
                    "flag": p[3].upper(), "owner": p[4].upper(), "ship_type": p[5].upper()
                })
                update_draft_calculations(vessel_data, vessel_data["draft_fore"], vessel_data["draft_aft"])

        if target != "act_dates" and old_act:
            vessel_data['act_number'] = old_act

        vessel_data['step'] = 'review'
        show_review_menu(chat_id)
    except Exception as e:
        log_error("Ошибка при разборе ручного текстового редактирования данных", e)

@bot.message_handler(func=lambda message: user_states.get(message.chat.id, {}).get('step') == 'waiting_for_start_num')
def handle_start_num_input(message):
    chat_id = message.chat.id
    try:
        try:
            start_num = int(message.text.strip())
        except:
            start_num = 1
        msg_to_edit = user_states[chat_id].get('msg_to_edit')
        wb = Workbook()
        ws = wb.active
        ws.append(["vessel_name", "wsa", "propeller_area"])
        ws.append([
            user_states[chat_id].get("vessel_name"),
            user_states[chat_id].get("wetted_surface"),
            user_states[chat_id].get("propeller_area")
        ])
        try:
            wb.save(os.path.join(BASE_DIR, "vessels_data.xlsx"))
        except:
            pass
        generate_all_documents(chat_id, start_num, msg_to_edit)
    except Exception as e:
        log_error("Ошибка при фиксации стартового номера документов", e)

# =====================================================================
# FASTAPI API СЕРВЕР ДЛЯ РАБОТЫ WEB APP И АНИМАЦИИ
# =====================================================================
generation_tasks = {}

def update_task_status(task_id, status, progress, message, error=None):
    if task_id:
        existing = generation_tasks.get(task_id, {})
        generation_tasks[task_id] = {
            **existing,
            "status": status,
            "progress": progress,
            "message": message,
            "error": error,
            "updated_at": time.time()
        }

class GenerateRequest(BaseModel):
    vessel_name: str
    ship_type: str
    imo: str
    flag: str
    owner: str
    master_name: str
    loa: str
    breadth: str
    draft_fore: str
    draft_aft: str
    act_number: str
    diving_date: str
    date_start: str
    date_end: str
    current_date: str
    wetted_surface: str
    propeller_area: str
    company_key: str
    start_num: int
    email_dest: str
    chat_id: int

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def bg_generate_task(task_id: str, req_data: GenerateRequest):
    try:
        chat_id = req_data.chat_id
        
        company_key = req_data.company_key
        co = COMPANIES_DATA.get(company_key, {})
        
        vessel_data = {
            "vessel_name":      req_data.vessel_name.upper(),
            "ship_type":        req_data.ship_type.upper(),
            "imo":              req_data.imo,
            "flag":             req_data.flag.upper(),
            "owner":            req_data.owner.upper(),
            "master_name":      req_data.master_name.upper(),
            "loa":              req_data.loa,
            "breadth":          req_data.breadth,
            "draft_fore":       req_data.draft_fore,
            "draft_aft":        req_data.draft_aft,
            "act_number":       req_data.act_number,
            "diving_date":      req_data.diving_date,
            "date_start":       req_data.date_start,
            "date_end":         req_data.date_end,
            "current_date":     req_data.current_date,
            "wetted_surface":   req_data.wetted_surface,
            "propeller_area":   req_data.propeller_area,
            "company_contract": co.get("company_contract", ""),
            "company_short":    co.get("company_short", "НЕ ВЫБРАНА"),
            "company_full":     co.get("company_full", ""),
            "email_dest":       req_data.email_dest,
        }
        
        user_states[chat_id] = vessel_data
        update_draft_calculations(user_states[chat_id], user_states[chat_id]["draft_fore"], user_states[chat_id]["draft_aft"])
        
        wb = Workbook()
        ws = wb.active
        ws.append(["vessel_name", "wsa", "propeller_area"])
        ws.append([
            vessel_data.get("vessel_name"),
            vessel_data.get("wetted_surface"),
            vessel_data.get("propeller_area")
        ])
        try:
            wb.save(os.path.join(BASE_DIR, "vessels_data.xlsx"))
        except:
            pass
            
        generate_all_documents(chat_id, req_data.start_num, msg_to_edit=None, task_id=task_id)
        
    except Exception as e:
        log_error(f"API bg_generate_task error: {e}", e)
        update_task_status(task_id, "failed", 100, f"Ошибка генерации: {str(e)}", error=str(e))

@app.post("/api/generate")
async def api_generate(req: GenerateRequest, background_tasks: BackgroundTasks):
    task_id = str(uuid.uuid4())
    generation_tasks[task_id] = {
        "status": "preparing",
        "progress": 5,
        "message": "Инициализация...",
        "error": None,
        "updated_at": time.time(),
        "chat_id": req.chat_id
    }
    background_tasks.add_task(bg_generate_task, task_id, req)
    return {"task_id": task_id}

@app.get("/api/status/{task_id}")
async def api_status(task_id: str):
    task = generation_tasks.get(task_id)
    if not task:
        return {"status": "not_found", "progress": 0, "message": "Задача не найдена"}
    return task

@app.get("/api/history")
async def api_get_history():
    db = load_history_db()
    return db.get("used_acts", [])

@app.post("/api/history/delete")
async def api_delete_history(act_number: str):
    db = load_history_db()
    db["used_acts"] = [a for a in db.get("used_acts", []) if a.get("act_number") != act_number]
    save_history_db(db)
    return {"status": "success"}

@app.get("/api/system/status")
async def api_system_status():
    return {"bot_active": BOT_ACTIVE}

@app.get("/api/debug/webapp-url")
async def api_debug_webapp_url():
    return {"webapp_url": WEBAPP_URL}

@app.get("/api/debug/env")
async def api_debug_env():
    return {
        "TELEGRAM_BOT_TOKEN_set": bool(TELEGRAM_BOT_TOKEN),
        "MY_CHAT_ID": MY_CHAT_ID,
        "SECOND_BOT_TOKEN_len": len(SECOND_BOT_TOKEN) if SECOND_BOT_TOKEN else 0,
        "SECOND_BOT_TOKEN_mask": f"{SECOND_BOT_TOKEN[:10]}...{SECOND_BOT_TOKEN[-5:]}" if SECOND_BOT_TOKEN else "None",
        "SECOND_BOT_CHAT_ID": SECOND_BOT_CHAT_ID,
        "EMAIL_SENDER": EMAIL_SENDER,
        "PROXY_URL_set": bool(PROXY_URL)
    }

@app.get("/api/debug/errors")
async def api_debug_errors():
    log_path = os.path.join(BASE_DIR, "terminator_errors.log")
    if not os.path.exists(log_path):
        return {"message": "No errors log file found"}
    try:
        with open(log_path, "r", encoding="utf-8") as f:
            content = f.read()
        return {"log": content[-20000:]}
    except Exception as e:
        return {"error": str(e)}

@app.post("/api/upload-pdf")
async def api_upload_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        return JSONResponse(status_code=400, content={"error": "Только файлы PDF!"})
    
    temp_filename = f"upload_{uuid.uuid4()}_{file.filename}"
    temp_path = os.path.join(BASE_DIR, temp_filename)
    try:
        content = await file.read()
        with open(temp_path, "wb") as f:
            f.write(content)
            
        data_ai = analyze_combined_email_data("", [temp_path])
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
        if not data_ai:
            return JSONResponse(status_code=422, content={"error": "Не удалось распознать данные судна ИИ-парсером."})
            
        div_date = data_ai.get("diving_date", datetime.now().strftime("%d.%m.%Y"))
        if "DD.MM.YYYY" in div_date or not div_date:
            div_date = datetime.now().strftime("%d.%m.%Y")
            
        try:
            dt_obj = datetime.strptime(div_date.strip(), "%d.%m.%Y")
            date_start_str = (dt_obj - timedelta(days=1)).strftime("%d.%m.%Y")
            date_end_str = (dt_obj + timedelta(days=3)).strftime("%d.%m.%Y")
        except:
            date_start_str = datetime.now().strftime("%d.%m.%Y")
            date_end_str = (datetime.now() + timedelta(days=4)).strftime("%d.%m.%Y")
            
        result = {
            "vessel_name":  str(data_ai.get("vessel_name", "UNKNOWN")).upper(),
            "ship_type":    str(data_ai.get("ship_type", "GENERAL CARGO")).upper(),
            "imo":          str(data_ai.get("imo", "0000000")),
            "loa":          str(data_ai.get("loa", "0.0")),
            "breadth":      str(data_ai.get("breadth", "0.0")),
            "flag":         str(data_ai.get("flag", "НЕ УКАЗАН")).upper(),
            "owner":        str(data_ai.get("owner", "НЕ УКАЗАН")).upper(),
            "draft_fore":   str(data_ai.get("draft_fore", "0.0")),
            "draft_aft":    str(data_ai.get("draft_aft", "0.0")),
            "master_name":  str(data_ai.get("master_name", "НЕ НАЙДЕН")).upper(),
            "act_number":   get_next_act_number_for_date(div_date),
            "diving_date":  div_date,
            "date_start":   date_start_str,
            "date_end":     date_end_str,
            "current_date": datetime.now().strftime("%d.%m.%Y"),
            "company_key":  "",
            "last_doc_number": load_history_db().get("last_doc_number", 145)
        }
        return result
    except Exception as e:
        log_error("Ошибка при асинхронной загрузке PDF через API", e)
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except: pass
        return JSONResponse(status_code=500, content={"error": f"Внутренняя ошибка сервера: {str(e)}"})

@app.get("/api/download/zip/{task_id}")
async def api_download_zip(task_id: str):
    task = generation_tasks.get(task_id)
    if not task:
        return JSONResponse(status_code=404, content={"error": "Задача не найдена"})
        
    zip_path = task.get("zip_path")
    vessel_name = task.get("vessel_name", "archive")
    if not zip_path or not os.path.exists(zip_path):
        return JSONResponse(status_code=404, content={"error": "Файл архива не найден или был удален"})
        
    return FileResponse(
        path=zip_path,
        media_type="application/zip",
        filename=f"DOCUMENTS_{vessel_name}.zip"
    )

@app.get("/api/download/file/{task_id}/{file_name}")
async def api_download_file(task_id: str, file_name: str):
    task = generation_tasks.get(task_id)
    if not task:
        return JSONResponse(status_code=404, content={"error": "Задача не найдена"})
        
    vessel_folder = task.get("vessel_folder")
    if not vessel_folder or not os.path.exists(vessel_folder):
        return JSONResponse(status_code=404, content={"error": "Папка с документами не найдена"})
        
    clean_name = os.path.basename(file_name)
    file_path = os.path.join(vessel_folder, clean_name)
    
    if not os.path.exists(file_path):
        return JSONResponse(status_code=404, content={"error": "Запрошенный файл не найден"})
        
    return FileResponse(
        path=file_path,
        filename=clean_name
    )

@app.post("/api/system/toggle")
async def api_system_toggle():
    global BOT_ACTIVE
    BOT_ACTIVE = not BOT_ACTIVE
    return {"status": "success", "bot_active": BOT_ACTIVE}

@app.post("/api/system/restart")
async def api_system_restart(background_tasks: BackgroundTasks):
    def exit_after_delay():
        time.sleep(1)
        os._exit(0)
    background_tasks.add_task(exit_after_delay)
    return {"status": "success", "message": "Бот перезагружается..."}

class ConfirmRequest(BaseModel):
    action: str

@app.post("/api/confirm/{task_id}")
async def api_confirm(task_id: str, req: ConfirmRequest):
    task = generation_tasks.get(task_id)
    if not task:
        return {"status": "error", "message": "Задача не найдена"}
        
    chat_id = task.get("chat_id")
    vessel_data = user_states.get(chat_id)
    
    if req.action == "send":
        update_task_status(task_id, "sending_emails", 90, "Отправка писем...")
        
        email_dest = vessel_data.get("email_dest") if vessel_data else ""
        if not email_dest:
            update_task_status(task_id, "failed", 100, "Ошибка: Email не задан")
            return {"status": "error", "message": "Email не задан"}
            
        generated = vessel_data.get("generated_files", [])
        if not generated:
            update_task_status(task_id, "failed", 100, "Ошибка: Нет файлов для отправки")
            return {"status": "error", "message": "Файлы не найдены"}
            
        # 1. Первый пакет (PDF документы)
        pdfs = [p for p in generated if p.lower().endswith('.pdf')]
        files_to_send1 = pdfs if pdfs else generated[:3]
        success1 = send_email_with_attachments(email_dest, files_to_send1, vessel_data.get("vessel_name", "Vessel"))
        
        # 2. Второй пакет (первые 3 DOCX документа - Акты)
        docx_files = generated[:3]
        success2 = False
        if docx_files:
            success2 = send_email_with_attachments(email_dest, docx_files, vessel_data.get("vessel_name", "Vessel"))
            
        if success1 and success2:
            update_task_status(task_id, "completed", 100, "Все документы успешно отправлены!")
            bot.send_message(chat_id, f"✅ Оба пакета успешно отправлены на `{email_dest}`!", parse_mode="Markdown")
            return generation_tasks[task_id]
        else:
            msg = "⚠️ Возникли проблемы при отправке:\n"
            msg += f"Пакет 1 (PDF): {'Отправлен' if success1 else 'Ошибка'}\n"
            msg += f"Пакет 2 (DOCX): {'Отправлен' if success2 else 'Ошибка'}"
            update_task_status(task_id, "failed", 100, f"Ошибка SMTP: {msg}")
            bot.send_message(chat_id, msg, parse_mode="Markdown")
            return {"status": "partial_success", "message": msg}
            
    elif req.action == "cancel":
        update_task_status(task_id, "completed", 100, "Архив готов, отправка почты пропущена.")
        if chat_id in user_states:
            del user_states[chat_id]
        bot.send_message(chat_id, "🔒 Отправка почты отменена. Готовый архив доступен для скачивания в приложении.")
        return generation_tasks[task_id]

# =====================================================================
# ФОНОВЫЙ ПОТОК POLLING — ручной цикл, совместим с прокси
# =====================================================================
def polling_worker():
    """Ручной polling loop — short polling каждые 2с, не зависит от infinity_polling."""
    ALLOWED_UPDATES = [
        "message", "edited_message", "callback_query",
        "inline_query", "web_app_data"
    ]
    last_update_id = 0
    poll_fail_count = 0

    # Пропускаем накопившиеся старые сообщения
    try:
        initial = bot.get_updates(offset=-1, timeout=0, allowed_updates=ALLOWED_UPDATES)
        if initial:
            last_update_id = initial[-1].update_id
    except Exception as e:
        log_error("Не удалось получить начальный offset", e)

    while True:
        if not BOT_ACTIVE:
            time.sleep(5)
            continue
        try:
            updates = bot.get_updates(
                offset=last_update_id + 1,
                timeout=0,              # short polling — прокси не режет короткие запросы
                allowed_updates=ALLOWED_UPDATES
            )
            poll_fail_count = 0
            if updates:
                bot.process_new_updates(updates)
                last_update_id = updates[-1].update_id
            time.sleep(2)
        except Exception as e:
            poll_fail_count += 1
            delay = min(5 * poll_fail_count, 60)
            log_error(f"Ошибка polling (попытка {poll_fail_count}), пауза {delay}с", e)
            time.sleep(delay)

# =====================================================================
# ГЛАВНЫЙ ЗАПУСК — все сервисы в фоновых потоках, главный держит процесс
# =====================================================================
if __name__ == '__main__':
    _diag = open(os.path.join(BASE_DIR, "bot_diag.log"), "a", encoding="utf-8")
    def _p(msg):
        ts = datetime.now().strftime("%H:%M:%S")
        _diag.write(f"[{ts}] {msg}\n"); _diag.flush()

    _p("=== СТАРТ ===")
    threading.Thread(target=mail_checker_daemon, daemon=True).start()
    _p("mail_checker_daemon запущен")
    threading.Thread(target=polling_worker, daemon=True).start()
    _p("polling_worker запущен")
    _p("Бот работает. Управление через Telegram (/status, /acts)")

    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    _p(f"FastAPI + Uvicorn запускаются на порту {port}")
    uvicorn.run("webterminator:app", host="0.0.0.0", port=port, reload=False, log_config=None)
